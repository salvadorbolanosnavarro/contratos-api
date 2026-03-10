"""
generar_contratos.py
Genera contratos de arrendamiento y promesa de compraventa
a partir de los machotes reales, produciendo .docx + .pdf

Uso:
  python generar_contratos.py arrendamiento datos.json
  python generar_contratos.py compraventa datos.json
"""

import sys
import json
import copy
import subprocess
import os
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# ── Rutas de machotes ─────────────────────────────────────────────────────────
MACHOTE_ARRENDAMIENTO = "/mnt/user-data/uploads/CONTRATO_ARRENDAMIENTO_MACHOTE__2_.docx"
MACHOTE_COMPRAVENTA   = "/mnt/user-data/uploads/CORRECIONES_COMPRAVENTA_BUENO.docx"


# ─────────────────────────────────────────────────────────────────────────────
# UTILIDADES DE REEMPLAZO
# Trabaja directo en el XML para preservar formato original 100%
# ─────────────────────────────────────────────────────────────────────────────

def reemplazar_en_runs(parrafo, variables: dict):
    """
    Reemplaza variables en un párrafo manejando el caso donde
    una variable está partida entre varios runs (comportamiento normal de Word).
    """
    # Reconstruir texto completo del párrafo
    texto_completo = "".join(run.text for run in parrafo.runs)

    # Verificar si hay alguna variable en el texto completo
    hubo_cambio = False
    for clave, valor in variables.items():
        if clave in texto_completo:
            texto_completo = texto_completo.replace(clave, str(valor))
            hubo_cambio = True

    if not hubo_cambio:
        return

    # Si hubo cambio, poner todo el texto en el primer run y vaciar los demás
    # preservando el formato del primer run
    if parrafo.runs:
        parrafo.runs[0].text = texto_completo
        for run in parrafo.runs[1:]:
            run.text = ""


def reemplazar_en_documento(doc: Document, variables: dict):
    """Recorre todos los párrafos y tablas del documento reemplazando variables."""
    # Párrafos directos
    for parrafo in doc.paragraphs:
        reemplazar_en_runs(parrafo, variables)

    # Párrafos dentro de tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_en_runs(parrafo, variables)

    # Headers y footers
    for seccion in doc.sections:
        for parrafo in seccion.header.paragraphs:
            reemplazar_en_runs(parrafo, variables)
        for parrafo in seccion.footer.paragraphs:
            reemplazar_en_runs(parrafo, variables)


# ─────────────────────────────────────────────────────────────────────────────
# CONTRATO DE ARRENDAMIENTO
# ─────────────────────────────────────────────────────────────────────────────

def generar_arrendamiento(datos: dict, ruta_salida_docx: str, ruta_salida_pdf: str):
    """
    datos = {
        "fecha_firma":               "Morelia, Michoacán, a 15 de enero de 2026",
        "nombre_arrendador":         "JUAN PÉREZ GARCÍA",
        "nombre_arrendatario":       "MARÍA LÓPEZ MARTÍNEZ",
        "nombre_obligado_solidario": "CARLOS SÁNCHEZ RUIZ",
        "domicilio_inmueble":        "Calle Reforma No. 123, Col. Centro",
        "domicilio_arrendador":      "Av. Madero No. 456, Col. Chapultepec",
        "domicilio_obligado":        "Calle Juárez No. 789, Col. Jardines",
        "destino_uso":               "CASA HABITACIÓN",
        "fecha_inicio":              "01 de febrero de 2026",
        "fecha_terminacion":         "31 de enero de 2027",
        "monto_renta_numeros":       "$15,000.00",
        "monto_renta_letras":        "QUINCE MIL",
        "monto_deposito_numeros":    "$15,000.00",
        "monto_deposito_letras":     "QUINCE MIL",
        "forma_pago":                "transferencia bancaria a la cuenta CLABE 012345678901234567 del banco BBVA",
        "fecha_pago_renta":          "5 (cinco)",
        "fecha_nuevo_contrato":      "01 de febrero de 2027"
    }
    """
    doc = Document(MACHOTE_ARRENDAMIENTO)

    variables = {
        "[FECHA DE FIRMA DEL CONTRATO]":      datos["fecha_firma"],
        "[NOMBRE DEL ARRENDADOR]":            datos["nombre_arrendador"].upper(),
        "[NOMBRE DEL ARRENDATARIO]":          datos["nombre_arrendatario"].upper(),
        "[NOMBRE DEL OBLIGADO SOLIDARIO]":    datos["nombre_obligado_solidario"].upper(),
        "[DOMICILIO DEL INMUEBLE]":           datos["domicilio_inmueble"].upper(),
        "[DOMICILIO DEL ARRENDADOR]":         datos["domicilio_arrendador"].upper(),
        "[DOMICILIO DEL OBLIGADO SOLIDARIO]": datos["domicilio_obligado"].upper(),
        "[DESTINO Y USO DEL INMUEBLE]":       datos["destino_uso"].upper(),
        "[FECHA DE INICIO]":                  datos["fecha_inicio"],
        "[FECHA DE TERMINACIÓN]":             datos["fecha_terminacion"],
        "[MONTO DE RENTA EN NÚMEROS]":        datos["monto_renta_numeros"],
        "[MONTO DE RENTA EN LETRAS]":         datos["monto_renta_letras"].upper(),
        "[MONTO DEPÓSITO EN NÚMEROS]":        datos["monto_deposito_numeros"],
        "[MONTO DEPÓSITO EN LETRAS]":         datos["monto_deposito_letras"].upper(),
        "[FORMA DE PAGO]":                    datos["forma_pago"],
        "[FECHA DE PAGO DE RENTA]":           datos["fecha_pago_renta"],
        "[FECHA DEL NUEVO CONTRATO]":         datos["fecha_nuevo_contrato"],
    }

    reemplazar_en_documento(doc, variables)
    doc.save(ruta_salida_docx)
    print(f"✅ DOCX guardado: {ruta_salida_docx}")

    _convertir_a_pdf(ruta_salida_docx, ruta_salida_pdf)


# ─────────────────────────────────────────────────────────────────────────────
# PROMESA DE COMPRAVENTA
# ─────────────────────────────────────────────────────────────────────────────

def generar_compraventa(datos: dict, ruta_salida_docx: str, ruta_salida_pdf: str):
    """
    datos = {
        "fecha_contrato":              "07 de octubre de 2025",
        "nombre_vendedora":            "ANA TORRES MEDINA",
        "nombre_comprador":            "LUIS HERRERA SOTO",
        "domicilio_inmueble":          "Av. Acueducto No. 500",
        "colonia_inmueble":            "Chapultepec Norte",
        "cp_inmueble":                 "58260",
        "numero_escritura":            "12,345",
        "nombre_notario":              "LIC. ROBERTO SILVA MORA",
        "numero_notaria":              "15",
        "tomo_rpp":                    "245",
        "registro_rpp":                "6789",
        "domicilio_vendedora":         "Av. Acueducto No. 500, Col. Chapultepec Norte, CP 58260",
        "domicilio_comprador":         "Calle Pino No. 30, Col. Las Rosas, CP 58200",
        "precio_total_letras":         "TRES MILLONES QUINIENTOS MIL",
        "precio_total_numeros":        "$3,500,000.00",
        "monto_arras_letras":          "CIEN MIL",
        "monto_arras_numeros":         "$100,000.00",
        "monto_segundo_pago_letras":   "TRES MILLONES CUATROCIENTOS MIL",
        "monto_segundo_pago_numeros":  "$3,400,000.00",
        "cuenta_bancaria_vendedora":   "012345678901234567",
        "banco_vendedora":             "BBVA",
        "fecha_limite_segundo_pago":   "09 de octubre de 2025",
        "fecha_limite_escritura":      "30 de noviembre de 2025",
        "pena_convencional_comprador_letras":  "CIEN MIL",
        "pena_convencional_comprador_numeros": "$100,000.00",
        "pena_convencional_vendedora_letras":  "CIEN MIL",
        "pena_convencional_vendedora_numeros": "$100,000.00"
    }
    """
    doc = Document(MACHOTE_COMPRAVENTA)

    # El machote usa "A" como placeholder genérico en distintos contextos
    # Reemplazamos campo por campo con find/replace contextual
    # usando los textos exactos que encontramos en el machote

    variables = {
        # Fecha
        "07/10/2025":                   datos["fecha_contrato"],

        # Nombre vendedora — aparece como "C. A, PROPIETARIA" y "LA C. A"
        "LA C. A, PROPIETARIA":         f'LA C. {datos["nombre_vendedora"].upper()}, PROPIETARIA',
        "LA C. A promete VENDER":       f'LA C. {datos["nombre_vendedora"].upper()} promete VENDER',
        "la C. A promete":              f'la C. {datos["nombre_vendedora"].upper()} promete',
        "C. A, PROPIETARIA":            f'C. {datos["nombre_vendedora"].upper()}, PROPIETARIA',

        # Nombre comprador
        "EL C. A, A QUIEN":             f'EL C. {datos["nombre_comprador"].upper()}, A QUIEN',
        "el C. A promete COMPRAR":      f'el C. {datos["nombre_comprador"].upper()} promete COMPRAR',
        "EL PROMITENTE COMPRADOR\"":    f'EL PROMITENTE COMPRADOR"',

        # Domicilio inmueble (aparece como "A A A, COLONIA A, CÓDIGO POSTAL A")
        "A A A, COLONIA A, CÓDIGO POSTAL A": (
            f'{datos["domicilio_inmueble"].upper()}, '
            f'COLONIA {datos["colonia_inmueble"].upper()}, '
            f'CÓDIGO POSTAL {datos["cp_inmueble"]}'
        ),

        # Escritura
        "escritura pública número  A pasada ante la fe de la Lic. A, notaria público número A": (
            f'escritura pública número {datos["numero_escritura"]} pasada ante la fe de la '
            f'{datos["nombre_notario"]}, notaria público número {datos["numero_notaria"]}'
        ),
        "tomo A y registro AAAA": f'tomo {datos["tomo_rpp"]} y registro {datos["registro_rpp"]}',

        # Domicilio vendedora para notificaciones
        "señala como domicilio para recibir cualquier tipo de notificación el ubicado en A A A, COLONIA A, CÓDIGO POSTAL A, CORRESPONDIENTE AL MUNICIPIO DE MORELIA, MICHOACÁN.": (
            f'señala como domicilio para recibir cualquier tipo de notificación el ubicado en '
            f'{datos["domicilio_vendedora"].upper()}, CORRESPONDIENTE AL MUNICIPIO DE MORELIA, MICHOACÁN.'
        ),

        # Domicilio comprador para notificaciones
        "señala como domicilio para recibir y oír notificaciones el ubicado en A A A, COLONIA A, CÓDIGO POSTAL A, CORRESPONDIENTE AL MUNICIPIO DE MORELIA, MICHOACÁN.": (
            f'señala como domicilio para recibir y oír notificaciones el ubicado en '
            f'{datos["domicilio_comprador"].upper()}, CORRESPONDIENTE AL MUNICIPIO DE MORELIA, MICHOACÁN.'
        ),

        # Precio total
        "precio pactado de  (A pesos 00/100 M.N.)": (
            f'precio pactado de {datos["precio_total_numeros"]} '
            f'({datos["precio_total_letras"].upper()} PESOS 00/100 M.N.)'
        ),

        # Arras
        "la cantidad de  (A PESOS 00/100 M.N.) en efectivo": (
            f'la cantidad de {datos["monto_arras_numeros"]} '
            f'({datos["monto_arras_letras"].upper()} PESOS 00/100 M.N.) en efectivo'
        ),

        # Segundo pago
        "La cantidad de  (A PESOS 00/100 M.N.) mediante transferencia a la cuenta no. A del banco A": (
            f'La cantidad de {datos["monto_segundo_pago_numeros"]} '
            f'({datos["monto_segundo_pago_letras"].upper()} PESOS 00/100 M.N.) '
            f'mediante transferencia a la cuenta no. {datos["cuenta_bancaria_vendedora"]} '
            f'del banco {datos["banco_vendedora"]}'
        ),

        # Fecha límite segundo pago
        "a más tardar el 09/10/2025.\n\n\tTERCERA": (
            f'a más tardar el {datos["fecha_limite_segundo_pago"]}.\n\n\tTERCERA'
        ),

        # Fecha límite escritura
        "se celebre a más tardar el 09/10/2025.": (
            f'se celebre a más tardar el {datos["fecha_limite_escritura"]}.'
        ),

        # Pena convencional comprador
        'pagará a LA PROMITENTE VENDEDORApor concepto de pena convencional, la cantidad de  (A PESOS 00/100 M.N.)': (
            f'pagará a LA PROMITENTE VENDEDORA por concepto de pena convencional, '
            f'la cantidad de {datos["pena_convencional_comprador_numeros"]} '
            f'({datos["pena_convencional_comprador_letras"].upper()} PESOS 00/100 M.N.)'
        ),

        # Pena convencional vendedora
        'deberán pagar a \"EL PROMITENTE COMPRADORla cantidad de  (A PESOS 00/100 M.N.)': (
            f'deberán pagar a "EL PROMITENTE COMPRADOR" la cantidad de '
            f'{datos["pena_convencional_vendedora_numeros"]} '
            f'({datos["pena_convencional_vendedora_letras"].upper()} PESOS 00/100 M.N.)'
        ),
    }

    reemplazar_en_documento(doc, variables)
    doc.save(ruta_salida_docx)
    print(f"✅ DOCX guardado: {ruta_salida_docx}")

    _convertir_a_pdf(ruta_salida_docx, ruta_salida_pdf)


# ─────────────────────────────────────────────────────────────────────────────
# CONVERSIÓN A PDF via LibreOffice
# ─────────────────────────────────────────────────────────────────────────────

def _convertir_a_pdf(ruta_docx: str, ruta_pdf: str):
    dir_salida = os.path.dirname(os.path.abspath(ruta_pdf))
    try:
        resultado = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", dir_salida, os.path.abspath(ruta_docx)],
            capture_output=True, text=True, timeout=60
        )
        # LibreOffice genera el PDF con el mismo nombre base
        nombre_base = os.path.splitext(os.path.basename(ruta_docx))[0]
        pdf_generado = os.path.join(dir_salida, nombre_base + ".pdf")
        if os.path.exists(pdf_generado) and pdf_generado != ruta_pdf:
            os.rename(pdf_generado, ruta_pdf)
        if os.path.exists(ruta_pdf):
            print(f"✅ PDF  guardado: {ruta_pdf}")
        else:
            print(f"⚠️  PDF no generado. stdout: {resultado.stdout} stderr: {resultado.stderr}")
    except Exception as e:
        print(f"⚠️  Error convirtiendo a PDF: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python generar_contratos.py [arrendamiento|compraventa] datos.json [prefijo_salida]")
        sys.exit(1)

    tipo   = sys.argv[1].lower()
    datos  = json.loads(open(sys.argv[2]).read())
    prefijo = sys.argv[3] if len(sys.argv) > 3 else tipo

    docx_out = f"{prefijo}.docx"
    pdf_out  = f"{prefijo}.pdf"

    if tipo == "arrendamiento":
        generar_arrendamiento(datos, docx_out, pdf_out)
    elif tipo == "compraventa":
        generar_compraventa(datos, docx_out, pdf_out)
    else:
        print(f"Tipo desconocido: {tipo}. Usa 'arrendamiento' o 'compraventa'.")
        sys.exit(1)
